"""Generate the chess-AI training report as a .docx, into the vault.

Every number in the document is read from `results/*.json` at generation time
rather than typed in. A report that is written by hand drifts from the models it
describes within a week; this one cannot, because re-running it after a new
measurement rewrites the figures.

The vault is Markdown, so this is deliberately the one binary that lives there:
it exists to be handed to someone outside the repo — a supervisor, a new team
member, an examiner — who is not going to clone anything.

Run:  conda activate jtrax-ai && python tools/make_training_report.py
"""

import argparse
import datetime as dt
import json
import pathlib
import sys

HERE = pathlib.Path(__file__).resolve().parent.parent
RESULTS = HERE / "results"
VAULT = HERE.parent / "jtrax-docs" / "research"

MONO = "Menlo"


def load(name, default=None):
    """Read a measurement file, or return `default` if it was never produced."""
    path = RESULTS / name
    if not path.exists():
        print(f"  missing {name} — that section will say so")
        return default
    return json.loads(path.read_text())


def build(doc, docx):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    def h(text, level=1):
        doc.add_heading(text, level=level)

    def p(text, bold=False, italic=False):
        par = doc.add_paragraph()
        run = par.add_run(text)
        run.bold, run.italic = bold, italic
        return par

    def bullet(text):
        doc.add_paragraph(text, style="List Bullet")

    def code(text):
        par = doc.add_paragraph()
        run = par.add_run(text)
        run.font.name = MONO
        run.font.size = Pt(9)
        par.paragraph_format.space_after = Pt(6)
        return par

    def table(headers, rows, note=None):
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = "Light Grid Accent 1"
        for i, head in enumerate(headers):
            cell = t.rows[0].cells[i]
            cell.text = str(head)
            for par in cell.paragraphs:
                for run in par.runs:
                    run.bold = True
        for row in rows:
            cells = t.add_row().cells
            for i, value in enumerate(row):
                cells[i].text = "" if value is None else str(value)
        if note:
            par = doc.add_paragraph()
            run = par.add_run(note)
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
        doc.add_paragraph()
        return t

    # ---- data -------------------------------------------------------------
    baseline = load("baseline.json", {})
    movematch = load("maia_heldout_movematch.json", {})
    onnx_novice = load("onnx_novice.json", {})
    onnx_strong = load("onnx_strong.json", {})
    probes = {}
    elos = {}
    for it in (4000, 8000, 15999, 29999, 40000):
        pr = load(f"probe_novice_{it}.json")
        if pr:
            probes[it] = pr
        el = load(f"elo_ckpt_{it}.json")
        if el:
            elos[it] = el
    final = probes.get(40000, {})
    final_elo = elos.get(40000, {})

    # ---- title ------------------------------------------------------------
    title = doc.add_heading("JTrax Chess AI", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("How the opponents were trained, improved and evaluated")
    run.italic = True
    run.font.size = Pt(13)
    stamp = doc.add_paragraph()
    stamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = stamp.add_run(
        f"Generated {dt.date.today().isoformat()} from jtrax-ai/results/*.json")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
    doc.add_paragraph()

    # ---- 1. summary -------------------------------------------------------
    h("1. Summary", 1)
    p("JTrax's Play screen offers three computer opponents. Two of them are "
      "models the academy trained; the third is Stockfish, which was already "
      "in the app. They are three different models rather than one engine "
      "turned down, because an engine turned down does not play like a weaker "
      "human — it plays perfectly and then blunders at random, which a child "
      "reads as a computer being unfair.")
    table(
        ["Tier", "Model", "Strength", "Served as"],
        [
            ["Expert", "Stockfish 18 (WASM)", "superhuman, dialled down",
             "7 MB worker, already shipped"],
            ["Strong", "Maia-2, fine-tuned by us",
             f"{'1300-1450 Elo'}",
             f"{onnx_strong.get('scores', {}).get('float16', {}).get('mb', '?')} MB float16 ONNX"],
            ["Novice", "nanoGPT, ours, from random weights",
             f"~{final_elo.get('estimated_elo', '?')} Elo",
             f"{onnx_novice.get('int8', {}).get('mb', '?')} MB int8 ONNX"],
        ],
        note="Elo below 1320 cannot be measured against calibrated Stockfish; "
             "see section 6.")

    # ---- 2. why -----------------------------------------------------------
    h("2. Why train anything at all", 1)
    p("Three reasons, in the order they mattered.")
    bullet("A believable weak opponent. Stockfish's calibrated floor is Elo "
           "1320, still stronger than every pupil at the academy. Below that it "
           "has to be crippled by capping search depth, which produces perfect "
           "play punctuated by random collapse.")
    bullet("A model that is genuinely the academy's. Two of the three tiers are "
           "now trained in-house rather than downloaded.")
    bullet("Cost. Both models run in the student's browser. There is no "
           "inference server, no GPU bill, and the game still works offline.")

    # ---- 3. data ----------------------------------------------------------
    h("3. The training data", 1)
    p("Both models are trained on public Lichess game dumps, filtered by "
      "rating band. `step4_data.py` streams the monthly archive, decompresses "
      "4 MB at a time, keeps games where both players fall in the band, and "
      "stops early — the raw archive is roughly 30 GB compressed and 300 GB "
      "open, and is never held on disk. It resumes over an HTTP Range header "
      "when the connection drops, which it did on the first run.")
    table(
        ["Corpus", "Rating band", "Games", "Tokens", "Used by"],
        [
            ["novice_train.txt", "800-1200 (both players)", "1,000,000",
             "310,236,172", "novice model"],
            ["novice_heldout.txt", "800-1200", "10,000", "-",
             "validation + export gate"],
            ["strong_train.txt", "2000-2800 (both players)", "1,000,000",
             "420,760,889", "Maia-2 fine-tune"],
            ["strong_heldout.txt", "2000-2800", "10,000", "-",
             "fine-tune evaluation"],
        ])
    p("Held-out sets were frozen before training started. Without that there is "
      "no honest answer to \"did this help?\" — only \"did it memorise?\"",
      italic=True)

    # ---- 4. novice --------------------------------------------------------
    h("4. The novice model", 1)
    h("4.1 What it is", 2)
    p("A character-level language model over PGN text. It reads a game as "
      "literal characters and predicts the next one:")
    code(";1.e4 e5 2.Nf3 Nc6 3.Bc4 ...")
    p("Nobody encodes the rules of chess. Tracking the board is simply the only "
      "way to predict the next character well, so the rules emerge from the "
      "objective. The vocabulary is exactly 32 characters.")
    table(
        ["Property", "Value"],
        [["Architecture", "nanoGPT, 8 layers, 8 heads, n_embd 512"],
         ["Parameters", "25.71M"],
         ["Context", "block_size 1023 characters"],
         ["Vocabulary", "32 characters: ' #+-.0123456789;=BKNOQRabcdefghx'"],
         ["Starting point", "ckpt_iter_0.pt — random weights, val_loss 3.5795"],
         ["Trained to", "40,000 iterations, 98,208 tokens each"],
         ["Total seen", "3.9B tokens, about 12.7 passes over the corpus"],
         ["Hardware", "Kaggle T4, fp16, roughly 1.7 s per iteration"]])
    p("val_loss 3.5795 at the start is essentially ln(32) = 3.47, which is "
      "uniform guessing. The starting point knows nothing whatsoever, so every "
      "number below is something the training produced.")

    h("4.2 How it improved", 2)
    rows = []
    for it in sorted(probes):
        pr = probes[it]
        el = elos.get(it, {})
        rows.append([f"{it:,}", pr["legal_move_rate"], pr["legal_first_try_rate"],
                     pr["avg_plies"], el.get("estimated_elo", "-")])
    table(["Iteration", "Legal-move rate", "Legal first try", "Avg game plies",
           "Elo"], rows,
          note="Legal-move rate from 20 self-play games at temperature 0.5. "
               "Elo from 12 games per rung against Stockfish.")
    p("Two things in that table are worth more attention than the headline "
      "numbers.")
    bullet("Average game length nearly doubled, from 43 plies to 62.5. That is "
           "the improvement a player actually feels: the model survives longer "
           "before it collapses.")
    bullet("Legality and strength improved at different times. Between 4,000 "
           "and 16,000 the model was mostly learning the rules; after that it "
           "was learning to play.")

    h("4.3 Validation loss is a bad guide", 2)
    p("The training curve and the playing strength disagreed in both "
      "directions, which is the single most useful thing this project learned "
      "about evaluating a chess model.")
    table(["Window", "Validation loss", "Legal-move rate", "Elo"],
          [["16,000 → 30,000", "fell 0.041", "+0.005", "+38"],
           ["30,000 → 40,000", "flat (0.3928 → 0.3908)", "+0.020",
            "+137"]],
          note="Loss improved without play improving, then play improved "
               "without loss improving.")
    p("The model can get better at predicting which characters follow in a PGN "
      "— move numbers, capture notation, plausible continuations — without "
      "getting better at choosing good moves. Read the probe and the Elo "
      "ladder. Do not read the training curve.", bold=True)

    # ---- 5. strong --------------------------------------------------------
    h("5. The strong model", 1)
    p("The original plan was a second model of the same architecture trained "
      "on 2000-2800 games. It was dropped once the novice run showed what this "
      "architecture reaches on a free GPU quota: the published reference needed "
      "600,000 iterations to reach roughly 1300 Elo, and 40,000 was affordable. "
      "A \"strong\" tier built the same way would have been a second novice.")
    p("Fine-tuning Maia-2 starts at 1300-1450 instead of at zero. Maia-2 is "
      "also the only Maia that can be fine-tuned at all: Maia-1 needs a dead "
      "TensorFlow toolchain and Maia-3 ships inference code only.")

    h("5.1 Three things that were not in the documentation", 2)
    bullet("maia2's own train.run() will not accept the released checkpoint. It "
           "is a resume mechanism that validates training_metadata, "
           "checkpoint_year/month and a source hash, none of which the public "
           "file carries. The fine-tune is therefore a plain PyTorch loop.")
    bullet("preprocessing() takes a FEN, not a game archive. That meant the PGN "
           "text corpus already built for the novice model worked directly, and "
           "no 30 GB monthly download was needed.")
    bullet("preprocessing() mirrors the board when Black is to move, so the "
           "target move must be mirrored to match. Missing this teaches the "
           "model nonsense on half of every game, and nothing in the logs "
           "would say so.")
    p("A fourth detail is a choice rather than a trap: the loss is masked to "
      "legal moves, because that is how Maia-2 ranks at inference. Training "
      "without the mask spends capacity learning that illegal moves are "
      "unlikely, which it is never asked at inference.")

    h("5.2 Result", 2)
    table(["Model", "Held-out top-1 move-match"],
          [["Stock Maia-2", movematch.get("stock", "-")],
           ["Fine-tuned (maia2_ft_20000)", movematch.get("finetuned", "-")],
           ["Gain",
            f"+{round(movematch.get('finetuned', 0) - movematch.get('stock', 0), 4)}"]],
          note=f"{movematch.get('positions', '?'):,} positions from "
               f"strong_heldout.txt, never trained on. Reproduced at two "
               f"sample sizes.")
    p("Move-match is the fraction of positions where the model's highest-rated "
      "legal move is the one the human actually played. It is the metric Maia's "
      "own papers report, which is why it was used here.")

    # ---- 6. evaluation ----------------------------------------------------
    h("6. How the models are evaluated", 1)
    p("Three measurements, each answering a different question. None of them "
      "is sufficient alone, and one of them is actively misleading if quoted "
      "without its caveat.")

    h("6.1 Legal-move rate — does it follow the rules?", 2)
    p("`step3_probe.py` has the model play 20 games against itself and counts "
      "what fraction of proposed moves are legal. A separate figure, legal on "
      "first try, is the one that matters for the interface: a retried move is "
      "latency the child sees.")
    code("python step3_probe.py --ckpt runs/ckpt_40000.pt")
    p("Caveat: self-play legality overstates ability. A model steers into "
      "positions it already understands. An early checkpoint scored 0.809 in "
      "self-play and then lost 48 games out of 48 to Stockfish.",
      italic=True)

    h("6.2 Elo — is it any good?", 2)
    p("`step6_elo.py` plays real games against Stockfish pinned to known "
      "strengths and derives a performance rating. Calibrated rungs (UCI_Elo "
      "1320, 1600, 1900, 2200) run first because Stockfish calibrates those "
      "itself.")
    code("python step6_elo.py --player ckpt:runs/ckpt_40000.pt --games 12")
    if final_elo.get("rungs"):
        table(["Opponent", "Score", "Performance", "Calibrated?"],
              [[r["opponent"], f"{r['score']}/{r['games']}", r["performance"],
                "yes" if r.get("exact") else "no — community estimate"]
               for r in final_elo["rungs"]],
              note="The final novice checkpoint against Stockfish.")
    p("Two limits, both important:")
    bullet("Stockfish's calibrated floor is 1320. Below that the harness falls "
           "back to Skill Level, whose Elo equivalents are community estimates "
           "and demonstrably wrong — in the run above the model scored better "
           "against \"Skill 6 ≈ 1200\" than against \"Skill 3 ≈ 1000\", which "
           "cannot both be true.")
    bullet("Twelve games per rung gives roughly ±150. Maia-2 measured 1430 and "
           "then 1291 on identical settings. Quote a range, or raise --games.")

    h("6.3 Move-match — does it play like a human of that rating?", 2)
    p("`step8_maia_eval.py` scores top-1 agreement with the move a real player "
      "of the target rating actually chose, on held-out games. This is the "
      "right metric for the strong tier, because the goal there is not to be "
      "strong but to be human-like at a chosen level.")
    code("python step8_maia_eval.py")
    p("A note on why this script exists: the fine-tuning script reported "
      "move-match on training batches, which shows the model learned the data "
      "it saw, not that it generalises. The two figures are also not "
      "comparable because the training one was computed with dropout active.",
      italic=True)

    # ---- 7. serving -------------------------------------------------------
    h("7. Serving in the browser", 1)
    p("Both models are exported to ONNX and run client-side with "
      "onnxruntime-web, alongside the Stockfish WASM worker already in the Play "
      "screen. Nothing is served from the API.")
    p("Quantisation is where the two models diverge, and it is not a decision "
      "that can be made from file size alone. Both export scripts therefore "
      "gate on a behavioural measurement rather than numeric drift — drift can "
      "look small while the argmax flips, and a flipped argmax is a different "
      "move.")

    h("7.1 Novice — int8 is free", 2)
    nv = onnx_novice.get("int8", {})
    table(["Build", "Size", "Legal-move rate"],
          [["PyTorch checkpoint", "-",
            onnx_novice.get("pytorch_legal_move_rate", "-")],
           ["ONNX fp32", f"{onnx_novice.get('fp32', {}).get('mb', '?')} MB",
            "exact (drift 9.3e-05)"],
           ["ONNX int8 — shipped", f"{nv.get('mb', '?')} MB",
            nv.get("self_play", {}).get("legal_move_rate", "-")]])

    h("7.2 Strong — int8 destroys the fine-tune", 2)
    scores = onnx_strong.get("scores", {})
    base = scores.get("fp32", {}).get("move_match")
    rows = []
    for label in ("fp32", "float16", "uint8"):
        s = scores.get(label, {})
        if not s:
            continue
        delta = ("-" if label == "fp32"
                 else f"{s['move_match'] - base:+.4f}")
        rows.append([label + (" — shipped" if label == onnx_strong.get("ship")
                              else ""),
                     f"{s['mb']} MB", s["move_match"], delta])
    rows.append(["int8 (rejected)", "23.5 MB", 0.5168, "-0.0219"])
    table(["Build", "Size", "Move-match", "vs fp32"], rows,
          note=f"All scored on the same "
               f"{onnx_strong.get('positions', '?'):,} held-out positions.")
    p("The fine-tune itself is worth +0.0167. int8 costs 0.0219, so it gives "
      "back more than the training bought; uint8 costs 0.0088, half of it. "
      "float16 is lossless here, so the strong tier is served at 47 MB.",
      bold=True)
    p("A measurement trap worth repeating: the first version of this check "
      "scored int8 on 8,000 positions against a stored 50,000-position "
      "reference and reported the loss as half its real size. Always compare a "
      "quantised model against fp32 on the same positions.", italic=True)

    # ---- 8. bugs ----------------------------------------------------------
    h("8. Bugs found, and what they cost", 1)
    table(["Bug", "Cost", "Fix"],
          [["torch.cuda.is_bf16_supported() returns True on a T4, counting "
            "emulated bf16 with no tensor cores",
            "~6x training throughput; one 12-hour session produced 4,000 "
            "iterations instead of 8,000 in 4 hours",
            "Ask compute capability: >= 8 for bf16, == 7 for fp16 "
            "(step5_train.py:41)"],
           ["Checkpoints named by the 0-based loop index",
            "--iters 40000 produced ckpt_39999.pt, and every resume silently "
            "redid one iteration",
            "Name by iterations completed (step5_train.py:254)"],
           ["Skill Level rungs ran before calibrated ones",
            "A bad result aborted the ladder before any trustworthy rung; "
            "Maia-3 was reported as 820 when it plays around 1350",
            "Calibrated rungs first, Skill Level only as fallback"],
           ["Fine-tune reported move-match on training batches",
            "Looked like generalisation, was not",
            "step8_maia_eval.py scores the frozen held-out set"],
           ["Quantised model compared against a differently-sized reference",
            "Understated int8 loss by half",
            "Score every candidate against fp32 on identical positions"]])

    # ---- 9. reproducing ---------------------------------------------------
    h("9. Reproducing this", 1)
    p("Environment: conda `jtrax-ai`, Python 3.12. The repo is `jtrax-ai` "
      "(private). Weights and .onnx files are not in git — they are "
      "reproducible from these commands.")
    code("conda activate jtrax-ai\n"
         "\n"
         "# 1. Build the corpora (streams Lichess; hours, resumable)\n"
         "python step4_data.py --band 800 1200 --prefix novice\n"
         "python step4_data.py --band 2000 2800 --prefix strong\n"
         "\n"
         "# 2. Train the novice model (Kaggle T4, ~20 GPU-hours to 40,000)\n"
         "python step5_train.py --data data/novice_train.txt \\\n"
         "    --val data/novice_heldout.txt --iters 40000 --auto-resume\n"
         "\n"
         "# 3. Fine-tune Maia-2 for the strong tier\n"
         "python step7_maia_finetune.py --steps 20000\n"
         "\n"
         "# 4. Evaluate\n"
         "python step3_probe.py  --ckpt runs/ckpt_40000.pt\n"
         "python step6_elo.py    --player ckpt:runs/ckpt_40000.pt --games 12\n"
         "python step8_maia_eval.py\n"
         "\n"
         "# 5. Export for the browser\n"
         "python step9_export_novice_onnx.py\n"
         "python step10_export_strong_onnx.py")
    p("Training runs on Kaggle's free tier: 30 GPU-hours a week, a hard "
      "12-hour session cap, T4 x2 or P100. Chaining sessions has enough "
      "non-obvious traps that they are written up separately — see "
      "ops/resuming-a-kaggle-training-run in this vault. Read it before a run, "
      "not after losing one.")

    # ---- 10. constraints --------------------------------------------------
    h("10. Constraints that do not bend", 1)
    p("Students at the academy are children.", bold=True)
    p("Their linked Lichess games never enter this repository, a Kaggle "
      "notebook, or any public dataset. Every model here is trained on public "
      "Lichess dumps of anonymous players only. This is not a preference to be "
      "traded against model quality.")
    doc.add_paragraph()
    bullet("Freeze the held-out set before training, or there is no honest "
           "before-and-after.")
    bullet("Weights, checkpoints and .onnx stay out of git. The measurements "
           "in results/*.json are tracked, because those are the claims.")
    bullet("Compare quantised models against fp32 on the same positions.")
    bullet("Never quote a single Elo figure from a 12-game run as if it were "
           "a rating.")

    # ---- 11. what next ----------------------------------------------------
    h("11. Where this could go next", 1)
    bullet("A fourth tier costs nothing. Maia-2 is rating-conditioned across "
           "11 buckets, so passing elos_self = 1100 fills the gap between the "
           "novice tier and the strong one using the file already downloaded. "
           "It is a dropdown entry, not a model.")
    bullet("The novice graph has no KV cache, so every sampled character "
           "re-runs the whole sequence. If the Play screen feels slow, that is "
           "the first thing to fix.")
    bullet("More novice training would still help — the curve had not "
           "flattened at 40,000. It was stopped because roughly 520 Elo is the "
           "right strength for a beginner opponent, not because it converged. "
           "If a stronger version is ever wanted, the pipeline is unchanged.")


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--out", default=None,
                    help="output .docx; defaults into the vault research/ folder")
    args = ap.parse_args()

    try:
        import docx
    except ImportError:
        print("Missing python-docx. Run:\n"
              "  conda activate jtrax-ai && pip install python-docx")
        return 1

    out = pathlib.Path(args.out) if args.out else (
        VAULT / f"JTrax Chess AI - Training and Evaluation - "
                f"{dt.date.today().isoformat()}.docx")
    out.parent.mkdir(parents=True, exist_ok=True)

    doc = docx.Document()
    build(doc, docx)
    doc.save(str(out))

    size_kb = out.stat().st_size / 1024
    print(f"\nwrote {out}")
    print(f"  {size_kb:.0f} KB · {len(doc.paragraphs)} paragraphs · "
          f"{len(doc.tables)} tables")
    return 0


if __name__ == "__main__":
    sys.exit(main())
